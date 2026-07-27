import os
import subprocess
import sys

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("python-docx not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def build_docx():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Document Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("ASSIGNMENT REPORT - WEEK 10 (TASK 1)")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1E, 0x4B, 0x27) # Dark Green
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Developing a Django Web Application for IoT Data in a Remote Server")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    # Spacing
    doc.add_paragraph()
    
    # Metadata Table
    table = doc.add_table(rows=5, cols=2)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    metadata = [
        ("Student Name", "Muhammad Ketsar Ali Abi Wahid"),
        ("Student ID (NIM)", "230401070204"),
        ("Class", "IF601"),
        ("Course Name", "Internet of Things (IoT)"),
        ("Lecturer", "Prof. Jong-Dae Park")
    ]
    
    for idx, (label, val) in enumerate(metadata):
        row = table.rows[idx]
        cell_lbl = row.cells[0]
        cell_val = row.cells[1]
        
        cell_lbl.text = label
        cell_val.text = val
        
        set_cell_background(cell_lbl, "F0F4F1")
        set_cell_margins(cell_lbl, top=80, bottom=80, left=120, right=120)
        set_cell_margins(cell_val, top=80, bottom=80, left=120, right=120)
        
        cell_lbl.paragraphs[0].runs[0].font.bold = True
        cell_lbl.paragraphs[0].runs[0].font.size = Pt(10)
        cell_val.paragraphs[0].runs[0].font.size = Pt(10)
        
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Directory for screenshots
    ss_dir = r"e:\Documents\Documents Ketsar\UNSIA\MATKUL\SEMESTER 6\Internet of Things (IoT) - Prof. Jong-Dae Park\PERTEMUAN 10\JAWABAN\screenshoot"
    
    # Define tasks, images, and descriptions
    tasks = [
        {
            "title": "Step 1: Database Migration & Starting Gunicorn Web Server",
            "image": "ss-1.png",
            "desc_en": "Connected to the server via PuTTY, activated the Python virtual environment (venv), and executed 'python manage.py migrate' to apply database migrations. Ran 'python manage.py collectstatic --noinput' to compile static assets. Finally, copied Gunicorn control scripts and successfully launched the WSGI web server using './start_mine.sh'.",
            "desc_id": "Melakukan koneksi ke server melalui PuTTY, mengaktifkan python virtual environment (venv), dan menjalankan 'python manage.py migrate' untuk memigrasikan database. Menjalankan perintah 'python manage.py collectstatic --noinput' untuk mengumpulkan file statis. Terakhir, menyalin skrip kontrol Gunicorn dan mengaktifkan WSGI web server dengan './start_mine.sh'."
        },
        {
            "title": "Step 2: Local Simulation Client (send_django_mock_data.py)",
            "image": "ss-2.png",
            "desc_en": "Executed the custom Python automation script 'send_django_mock_data.py' locally. The script successfully formulated JSON payloads for temperature and humidity, sent 10 HTTP POST requests to the Django REST API endpoint at 'https://m70204.belajarhub.id/django/api/sensor/', and received success status responses from the server.",
            "desc_id": "Mengeksekusi skrip python otomatis 'send_django_mock_data.py' di terminal komputer lokal. Skrip berhasil menyusun payload JSON data suhu dan kelembapan, mengirimkan 10 HTTP POST request ke REST API Django di 'https://m70204.belajarhub.id/django/api/sensor/', dan mendapatkan respon status sukses dari server."
        },
        {
            "title": "Step 3: Tabular Sensor Log Verification (device_data_view)",
            "image": "ss-3.png",
            "desc_en": "Accessed the Django web application at 'https://m70204.belajarhub.id/django/device/device01/'. The page successfully queried the SQLite database, fetched the 10 recently sent sensor readings, and displayed them in an HTML table containing the correct timestamps, temperatures, and humidity percentages.",
            "desc_id": "Mengakses aplikasi web Django di 'https://m70204.belajarhub.id/django/device/device01/'. Halaman tersebut berhasil membaca database SQLite, mengambil 10 rekaman data sensor yang baru saja terkirim, dan menampilkannya pada tabel HTML lengkap dengan stempel waktu, suhu, dan persentase kelembapan."
        },
        {
            "title": "Step 4: Live Sensor Gauge Visualization (sensor_gauge_view)",
            "image": "ss-4.png",
            "desc_en": "Navigated to the visual gauge display at 'https://m70204.belajarhub.id/django/device/device01/gauge/'. The interface displays two semi-circular gauges representing the latest temperature (25.1 °C) and humidity (64.3 %) telemetry readings, providing real-time visual instrumentation.",
            "desc_id": "Membuka halaman visualisasi gauge di 'https://m70204.belajarhub.id/django/device/device01/gauge/'. Antarmuka menampilkan dua alat ukur setengah lingkaran (gauge) yang memvisualisasikan data telemetri terbaru, yaitu suhu (25.1 °C) dan kelembapan (64.3 %)."
        }
    ]
    
    for t_idx, task in enumerate(tasks):
        h = doc.add_paragraph()
        run_h = h.add_run(task["title"])
        run_h.font.name = 'Arial'
        run_h.font.size = Pt(13)
        run_h.font.bold = True
        run_h.font.color.rgb = RGBColor(0x1E, 0x4B, 0x27)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        
        # Insert Image
        img_path = os.path.join(ss_dir, task["image"])
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_after = Pt(6)
            run_img = p_img.add_run()
            run_img.add_picture(img_path, width=Inches(5.5))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = p_cap.add_run(f"Figure 1.{t_idx+1}. Screen capture demonstrating {task['title']}")
            run_cap.font.name = 'Arial'
            run_cap.font.size = Pt(9.5)
            run_cap.font.italic = True
            run_cap.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
            p_cap.paragraph_format.space_after = Pt(10)
        else:
            p_err = doc.add_paragraph()
            run_err = p_err.add_run(f"[PENTING: Masukkan gambar {task['image']} ke folder 'screenshoot']")
            run_err.font.color.rgb = RGBColor(0xFF, 0, 0)
            run_err.font.bold = True
            
        # Description
        p_desc = doc.add_paragraph()
        run_desc_lbl = p_desc.add_run("English Explanation:\n")
        run_desc_lbl.font.bold = True
        run_desc_lbl.font.size = Pt(10.5)
        run_desc_val = p_desc.add_run(task["desc_en"] + "\n\n")
        run_desc_val.font.size = Pt(10.5)
        
        run_desc_id_lbl = p_desc.add_run("Penjelasan Bahasa Indonesia:\n")
        run_desc_id_lbl.font.bold = True
        run_desc_id_lbl.font.italic = True
        run_desc_id_lbl.font.size = Pt(10)
        run_desc_id_lbl.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        run_desc_id_val = p_desc.add_run(task["desc_id"])
        run_desc_id_val.font.italic = True
        run_desc_id_val.font.size = Pt(10)
        run_desc_id_val.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        p_desc.paragraph_format.space_after = Pt(18)
        p_desc.paragraph_format.line_spacing = 1.15
        
        if t_idx < len(tasks) - 1:
            doc.add_page_break()
            
    dest_path = r"e:\Documents\Documents Ketsar\UNSIA\MATKUL\SEMESTER 6\Internet of Things (IoT) - Prof. Jong-Dae Park\PERTEMUAN 10\JAWABAN\TASK1_230401070204_MUHAMMAD KETSAR ALI ABI WAHID.docx"
    doc.save(dest_path)
    print(f"Report generated successfully at: {dest_path}")

if __name__ == '__main__':
    build_docx()
